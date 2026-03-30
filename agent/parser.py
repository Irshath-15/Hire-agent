import fitz
import docx
import os
import json
import tempfile
import subprocess
import platform
from groq import Groq
from dotenv import load_dotenv
from PIL import Image
import io

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def ocr_image_with_tesseract(image: Image.Image) -> str:
    """Advanced OCR optimized for scanned documents with image preprocessing."""
    try:
        import pytesseract
        import cv2
        import numpy as np
        from PIL import ImageEnhance, ImageFilter

        # Set tesseract path based on OS (cached)
        if not hasattr(ocr_image_with_tesseract, '_tesseract_path'):
            system = platform.system()

            if system == "Windows":
                ocr_image_with_tesseract._tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            elif system == "Linux":
                # Try multiple possible paths (only once)
                possible_paths = [
                    '/usr/bin/tesseract',
                    '/usr/local/bin/tesseract',
                    'tesseract'  # In PATH
                ]

                tesseract_path = None
                for path in possible_paths:
                    try:
                        result = subprocess.run([path, '--version'],
                                              capture_output=True, text=True, timeout=2)
                        if result.returncode == 0:
                            tesseract_path = path
                            break
                    except:
                        continue

                ocr_image_with_tesseract._tesseract_path = tesseract_path or 'tesseract'
            else:
                ocr_image_with_tesseract._tesseract_path = 'tesseract'

        pytesseract.pytesseract.tesseract_cmd = ocr_image_with_tesseract._tesseract_path

        # Convert PIL to OpenCV format
        img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # Image preprocessing for scanned documents
        processed_images = []

        # Original image
        processed_images.append(img_cv)

        # Grayscale conversion
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        processed_images.append(gray)

        # Noise reduction
        denoised = cv2.medianBlur(gray, 3)
        processed_images.append(denoised)

        # Contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        processed_images.append(enhanced)

        # Thresholding
        _, thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        processed_images.append(thresh)

        # Morphological operations
        kernel = np.ones((1,1), np.uint8)
        morph = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        processed_images.append(morph)

        # Convert back to PIL for different scaling
        pil_images = []
        for img in processed_images:
            if len(img.shape) == 2:  # Grayscale
                pil_img = Image.fromarray(img)
            else:  # Color
                pil_img = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))

            # Create different scales
            pil_images.append(pil_img)  # Original size
            pil_images.append(pil_img.resize((int(pil_img.width * 1.5), int(pil_img.height * 1.5)), Image.LANCZOS))  # 1.5x scale
            pil_images.append(pil_img.resize((int(pil_img.width * 2.0), int(pil_img.height * 2.0)), Image.LANCZOS))  # 2x scale

        # OCR configurations optimized for scanned documents
        ocr_configs = [
            '--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,-@()',
            '--psm 3 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,-@()',
            '--psm 1 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,-@()',
            '--psm 12 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,-@()',
            '--psm 4 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,-@()',
        ]

        best_text = ""
        best_confidence = 0

        # Try different combinations of preprocessing and OCR configs
        for pil_img in pil_images[:3]:  # Limit to first 3 preprocessed images for speed
            for config in ocr_configs[:2]:  # Limit to first 2 configs for speed
                try:
                    # Get both text and confidence data
                    data = pytesseract.image_to_data(pil_img, config=config, output_type=pytesseract.Output.DICT, timeout=10)

                    # Calculate average confidence
                    confidences = [int(conf) for conf in data['conf'] if conf != '-1']
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0

                    # Extract text
                    text = ' '.join([word for word in data['text'] if word.strip()])

                    # Keep the best result
                    if avg_confidence > best_confidence and text.strip():
                        best_text = text.strip()
                        best_confidence = avg_confidence

                        # If we have very good confidence, use it immediately
                        if avg_confidence > 70:
                            break

                except Exception as e:
                    continue

            if best_confidence > 70:
                break

        # If no good result found, try one more time with different preprocessing
        if best_confidence < 50 and len(best_text) < 50:
            try:
                # Convert to PIL and apply additional preprocessing
                pil_original = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))

                # Enhance contrast and sharpness
                enhancer = ImageEnhance.Contrast(pil_original)
                enhanced_img = enhancer.enhance(2.0)
                enhancer = ImageEnhance.Sharpness(enhanced_img)
                sharpened = enhancer.enhance(2.0)

                # Final OCR attempt
                final_data = pytesseract.image_to_data(sharpened, config='--psm 6', output_type=pytesseract.Output.DICT, timeout=15)
                final_text = ' '.join([word for word in final_data['text'] if word.strip()])

                if final_text.strip() and len(final_text) > len(best_text):
                    best_text = final_text.strip()

            except Exception as e:
                pass

        print(f"OCR completed with confidence: {best_confidence:.1f}%, text length: {len(best_text)}")
        return best_text

    except ImportError:
        print("OpenCV or pytesseract not available for advanced OCR")
        # Fallback to basic OCR
        try:
            import pytesseract
            pytesseract.pytesseract.tesseract_cmd = ocr_image_with_tesseract._tesseract_path
            result = pytesseract.image_to_string(image, config='--psm 6', timeout=15)
            return result.strip() if result else ""
        except:
            return ""
    except Exception as e:
        print(f"Advanced OCR failed: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> tuple:
    """Extract text from PDF. Optimized for speed with smart OCR."""
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        return f"[ERROR] Could not open PDF: {str(e)}", False

    text = ""
    has_any_content = False
    needs_ocr = False

    try:
        # Method 1: Quick standard text extraction - check first few pages
        sample_pages = min(3, len(doc))  # Check first 3 pages max

        for page_num in range(sample_pages):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text += page_text + "\n"
                has_any_content = True

        # If we got good text from sample, extract all pages quickly
        if has_any_content and len(text.strip()) > 100:
            for page_num in range(len(doc)):
                if page_num < sample_pages:  # Skip already processed pages
                    continue
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
            doc.close()
            return text.strip(), False

        # If sample pages had little/no text, check if document needs OCR
        needs_ocr = True

    except Exception as e:
        print(f"Standard extraction failed: {e}")
        needs_ocr = True

    # Method 2: Advanced OCR for scanned documents
    ocr_text = ""
    ocr_success = False

    if needs_ocr:
        print("Attempting advanced OCR extraction for scanned document...")
        try:
            # For scanned documents, process all pages with high quality
            max_pages = min(10, len(doc))  # Allow more pages for scanned docs

            for page_num in range(max_pages):
                try:
                    page = doc[page_num]

                    # Check if page has meaningful text (not just OCR artifacts)
                    existing_text = page.get_text()
                    word_count = len(existing_text.split()) if existing_text.strip() else 0

                    # If page has substantial text already, use it but also try OCR for comparison
                    if word_count > 20:
                        text += existing_text + "\n"
                        has_any_content = True
                        continue

                    # Render page to high-quality image for OCR
                    print(f"Processing page {page_num + 1}/{max_pages} with high-quality OCR...")
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)  # Higher resolution for scanned docs
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))

                    # Advanced OCR optimized for scanned documents
                    page_ocr = ocr_image_with_tesseract(img)

                    if page_ocr.strip() and len(page_ocr.split()) > 5:  # Require meaningful content
                        ocr_text += page_ocr + "\n"
                        ocr_success = True
                        has_any_content = True
                        print(f"✓ Page {page_num + 1}: extracted {len(page_ocr)} characters")
                    else:
                        print(f"⚠ Page {page_num + 1}: insufficient OCR text ({len(page_ocr) if page_ocr else 0} chars)")

                except Exception as page_err:
                    print(f"✗ Page {page_num + 1} OCR failed: {page_err}")
                    continue

        except Exception as e:
            print(f"OCR process failed: {e}")

    try:
        doc.close()
    except:
        pass

    # Combine and clean results
    final_text = ""
    if text.strip():
        final_text += text
    if ocr_text.strip():
        if final_text:
            final_text += "\n\n--- OCR Content ---\n\n"
        final_text += ocr_text

    if final_text.strip():
        print(f"✓ Total extracted text: {len(final_text)} characters from {max_pages if needs_ocr else sample_pages} pages")
        return final_text.strip(), ocr_success
    else:
        # Last resort: try all possible extraction methods
        print("Attempting final fallback extraction methods...")
        try:
            doc = fitz.open(file_path)
            fallback_text = ""

            for page in doc:
                # Try different extraction methods
                methods = [
                    lambda p: p.get_text("text"),
                    lambda p: p.get_text("blocks"),
                    lambda p: p.get_text("dict")
                ]

                for method in methods:
                    try:
                        page_text = method(page)
                        if isinstance(page_text, str) and page_text.strip():
                            fallback_text += page_text + "\n"
                            break
                        elif isinstance(page_text, list) and page_text:
                            # Handle block format
                            for block in page_text:
                                if isinstance(block, dict) and 'text' in block:
                                    fallback_text += block['text'] + "\n"
                            if fallback_text.strip():
                                break
                    except:
                        continue

            doc.close()
            if fallback_text.strip():
                print(f"✓ Fallback extraction successful: {len(fallback_text)} characters")
                return fallback_text.strip(), False

        except Exception as e:
            print(f"Fallback extraction failed: {e}")

        print("✗ All extraction methods failed")
        return "[ERROR] Could not extract any readable text from PDF - document may be corrupted, password-protected, or contain only images", True

def extract_text_from_docx(file_path: str) -> tuple:
    """Extract text from DOCX file - returns actual document text."""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text.strip(), False  # DOCX is never image-based
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return f"[ERROR] Could not read DOCX: {str(e)}", False

def extract_text(file_path: str) -> tuple:
    """Extract text from various file formats. Returns (text, is_image_based)."""
    ext = os.path.splitext(file_path)[1].lower()
    image_based = False
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        try:
            image = Image.open(file_path)
            ocr_text = ocr_image_with_tesseract(image).strip()
            if ocr_text:
                return ocr_text, True
            else:
                return f"[Image file: {os.path.basename(file_path)} -- no text detected]", True
        except Exception as e:
            return f"[Image file: {os.path.basename(file_path)} -- OCR failed: {str(e)}]", True
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def parse_resume_with_ai(raw_text: str) -> dict:
    """Parse resume with AI, optimized for speed."""
    # First try to extract name before AI processing
    inferred_name = infer_name_from_text(raw_text)

    # Truncate very long resumes for faster processing (keep first 4000 chars + last 1000)
    if len(raw_text) > 5000:
        truncated_text = raw_text[:4000] + "\n...\n" + raw_text[-1000:]
        print(f"Truncated resume from {len(raw_text)} to {len(truncated_text)} chars for faster processing")
    else:
        truncated_text = raw_text

    prompt = f"""You are a resume parser. Extract information from this resume and return ONLY a valid JSON object with no markdown, no code blocks, and no explanation. Return ONLY the JSON.

{{
  "name": "{inferred_name or 'null'}",
  "email": "email address or null",
  "phone": "phone number or null",
  "current_role": "most recent job title or null",
  "experience_years": "total years of experience as a number or null",
  "skills": "comma separated list of skills or null",
  "education": "highest qualification and institution or null",
  "red_flags": "any employment gaps, inconsistencies, vague roles — or null if none"
}}

IMPORTANT: Do NOT change or generate the name field. Keep it as: {inferred_name or 'null'}
Only extract email, phone, skills, experience, education, and red_flags from the resume text.

Resume text:
{truncated_text}"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=600,  # Reduced from 800
            timeout=10.0,    # Reduced from 15.0
            messages=[{"role": "user", "content": prompt}]
        )
    except Exception as e:
        print(f"Warning: Parser API timeout: {e}")
        # Return fallback with inferred data
        return {
            "name": inferred_name,
            "email": None,
            "phone": None,
            "current_role": None,
            "experience_years": None,
            "skills": None,
            "education": None,
            "red_flags": None
        }

    raw = response.choices[0].message.content
    
    # Clean the response: remove markdown code blocks
    clean = raw.replace("```json", "").replace("```", "").strip()
    
    # Try to extract JSON if it's wrapped in extra text
    if "{" in clean and "}" in clean:
        start = clean.find("{")
        end = clean.rfind("}") + 1
        clean = clean[start:end]

    try:
        result = json.loads(clean)
        # Ensure all required fields exist
        required_fields = ["name", "email", "phone", "current_role", "experience_years", "skills", "education", "red_flags"]
        for field in required_fields:
            value = result.get(field)
            # Convert string 'null'/'None' to Python None
            if value is None or value == 'null' or value == 'None' or value == '':
                result[field] = None
            else:
                result[field] = value
        
        # Force name to be the inferred one if available
        if inferred_name:
            result["name"] = inferred_name
        
        print(f"[PARSER] Parsed JSON successfully: name={result.get('name')}, email={result.get('email')}")
        return result
    except json.JSONDecodeError as e:
        print(f"[PARSER] JSON parse error: {e}, raw response: {raw[:200]}")
        return {
            "name": inferred_name,
            "email": None,
            "phone": None,
            "current_role": None,
            "experience_years": None,
            "skills": None,
            "education": None,
            "red_flags": None
        }

def infer_name_from_text(raw_text: str) -> str | None:
    """Try to infer candidate name from raw resume text - improved extraction."""
    import re

    if not raw_text or not raw_text.strip():
        return None

    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    
    # Strategy 1: Look for explicit "Name:" or "Contact:" patterns
    for line in lines:
        # Match patterns like "Name: John Doe" or "JOHN DOE"
        m = re.search(r"^\s*(?:name|candidate|applicant)\s*[:\-]\s*([A-Za-zÀ-ÖØ-öø-ÿ\s.,'-]+?)(?:\s*$|[,\|])", line, re.I)
        if m:
            candidate = m.group(1).strip()
            if candidate and 2 < len(candidate.split()) <= 5:
                return candidate
    
    # Strategy 2: Email prefix often contains name (PRIORITIZE THIS)
    email_match = re.search(r"([\w.%+-]+)@", raw_text)
    if email_match:
        local = email_match.group(1)
        # Convert email to name: john.doe → John Doe, john_silva → John Silva
        if '_' in local or '.' in local:
            name = re.sub(r'[._]', ' ', local).title()
            if name and len(name.split()) <= 5:
                print(f"[NAME_EXTRACT] Using email-derived name: '{name}'")
                return name
    
    # Strategy 3: First line that looks like a name (2-4 capitalized words)
    # But skip lines that are clearly not names
    for i, line in enumerate(lines[:15]):  # Check first 15 lines
        
        # Skip common headers and non-name content
        skip_keywords = ['resume', 'cv', 'curriculum', 'vitae', 'experience', 'summary', 'objective', 
                        'email', 'phone', 'address', 'linkedin', 'objective', 'skills', 'education',
                        'certifications', 'awards', 'projects', 'languages', 'level', 'contest',
                        'achievement', 'technical', 'additional', 'information', 'national', 'project',
                        'bachelor', 'technology', 'information', 'engineering', 'computer', 'science',
                        'university', 'college', 'school', 'institute', 'department', 'faculty']
        
        if any(skip in line.lower() for skip in skip_keywords):
            continue
        
        words = line.split()
        
        # Good candidate if: 1-4 words, all start with capital, all alphabetic, no numbers
        if 1 <= len(words) <= 4:
            if all(re.match(r"^[A-Za-zÀ-ÖØ-öø-ÿ'-]+$", w) for w in words):
                # Extra check: if it looks like a real name pattern
                if all(w[0].isupper() for w in words if w):
                    # Additional validation: not too long, no common non-name patterns
                    if len(line) <= 50 and not re.search(r'\d', line):  # No numbers
                        return line
    
    # Strategy 3.5: Look for initials + name pattern (S. DIVYA, etc.)
    for i, line in enumerate(lines[:10]):
        # Look for patterns like "S. DIVYA" or "JOHN DOE"
        if re.match(r'^[A-Z]\.\s+[A-Z]+', line):  # S. DIVYA
            return line
    
    # Strategy 4: Look for Phone/Email section header
    for i, line in enumerate(lines[:20]):
        if 'phone' in line.lower() or 'email' in line.lower():
            # Name is usually above contact info
            if i > 0:
                prev_line = lines[i-1]
                if len(prev_line.split()) <= 4 and all(w[0].isupper() for w in prev_line.split() if w):
                    return prev_line

    # Strategy 5: If no name found, try to derive from email as last resort
    email_match = re.search(r"([\w.%+-]+)@", raw_text)
    if email_match:
        local = email_match.group(1)
        # Try different variations
        variations = [
            local.replace('.', ' ').replace('_', ' ').title(),
            local.split('.')[0].title() if '.' in local else local.title(),
            local.split('_')[0].title() if '_' in local else local.title()
        ]
        for name in variations:
            if name and len(name.split()) <= 3:
                return name

    return None


def parse_resume(file_path: str) -> dict:
    try:
        raw_text, is_image_based = extract_text(file_path)
        print(f"[PARSER] Extracted {len(raw_text) if raw_text else 0} chars from {os.path.basename(file_path)}")
    except Exception as e:
        print(f"[PARSER] extract_text failed: {type(e).__name__}: {str(e)}")
        return {
            'name': None,
            'email': None,
            'phone': None,
            'current_role': None,
            'experience_years': None,
            'skills': None,
            'education': None,
            'red_flags': f'[ERROR] Text extraction failed: {str(e)}',
            'raw_text': '',
            'is_image_based': False
        }
    
    # If no text could be extracted, return appropriate error
    if not raw_text:
        print(f"[PARSER] No text extracted: empty")
        return {
            'name': None,
            'email': None,
            'phone': None,
            'current_role': None,
            'experience_years': None,
            'skills': None,
            'education': None,
            'red_flags': '[ERROR] Could not extract text - ensure file is readable',
            'raw_text': '',
            'is_image_based': False
        }
    
    # If text starts with error message, try to extract any usable text after it
    if raw_text.startswith('[') and 'ERROR' in raw_text:
        # Look for any actual content after the error message
        lines = raw_text.split('\n')
        actual_text = []
        for line in lines:
            if not line.startswith('[') or 'ERROR' not in line:
                actual_text.append(line)
        if actual_text:
            raw_text = '\n'.join(actual_text).strip()
            print(f"[PARSER] Stripped error prefix, using: {raw_text[:100]}...")
        else:
            # No usable text found, return error
            return {
                'name': None,
                'email': None,
                'phone': None,
                'current_role': None,
                'experience_years': None,
                'skills': None,
                'education': None,
                'red_flags': raw_text,
                'raw_text': raw_text,
                'is_image_based': is_image_based
            }

    # First, try to extract name BEFORE AI parsing
    try:
        inferred_name = infer_name_from_text(raw_text)
        print(f"[PARSER] Inferred name: {inferred_name}")
    except Exception as e:
        print(f"[PARSER] infer_name_from_text failed: {type(e).__name__}: {str(e)}")
        inferred_name = None
    
    # Then parse with AI - but use fallback if it fails
    try:
        parsed = parse_resume_with_ai(raw_text)
        print(f"[PARSER] AI parsing succeeded: {parsed.get('name')}")
    except Exception as e:
        print(f"[PARSER] parse_resume_with_ai failed: {type(e).__name__}: {str(e)}")
        # Use simple fallback parsing
        parsed = simple_extract(raw_text)
        parsed['name'] = inferred_name or parsed.get('name')
        print(f"[PARSER] Using fallback parser, name={parsed.get('name')}")
    
    parsed['raw_text'] = raw_text
    parsed['is_image_based'] = is_image_based

    # Ensure name is set
    if not parsed.get('name') and inferred_name:
        parsed['name'] = inferred_name
        print(f"[PARSER] Using inferred name: {inferred_name}")
    
    print(f"[PARSER] Final result: name={parsed.get('name')}, email={parsed.get('email')}")
    return parsed


def simple_extract(raw_text: str) -> dict:
    """Simple fallback extraction when AI parsing fails."""
    import re
    
    result = {
        'name': None,
        'email': None,
        'phone': None,
        'current_role': None,
        'experience_years': None,
        'skills': None,
        'education': None,
        'red_flags': None
    }
    
    try:
        # Try to find email
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
        if email_match:
            result['email'] = email_match.group(0)
        
        # Try to find phone
        phone_match = re.search(r'(?:\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}', raw_text)
        if phone_match:
            result['phone'] = phone_match.group(0)
        
        # Try to find name from first line or contact section
        lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
        if lines:
            first_line = lines[0]
            if not any(x in first_line.lower() for x in ['phone', 'email', 'address', 'linkedin']):
                result['name'] = first_line[:100]  # Use first line as name if it looks reasonable
        
        print(f"[PARSER] Simple extract: name={result.get('name')}, email={result.get('email')}")
    except Exception as e:
        print(f"[PARSER] Simple extract error: {e}")
    
    return result