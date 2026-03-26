#!/usr/bin/env python3
"""
Demo script to test resume parsing with a sample resume
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from agent.parser import parse_resume
from docx import Document

# Create a test resume using python-docx
test_resume_docx = "test_resume_sample.docx"

doc = Document()
doc.add_heading('JOHN SMITH', 0)
doc.add_paragraph('Email: john.smith@example.com | Phone: (555) 123-4567\n')

doc.add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph('Experienced Full-Stack Developer with 8 years of expertise in Python, JavaScript, and React. Proven track record building scalable web applications and leading engineering teams.\n')

doc.add_heading('EXPERIENCE', level=1)
doc.add_paragraph('Senior Software Engineer | TechCorp Inc. | 2021-Present', style='List Bullet')
doc.add_paragraph('Led development of microservices architecture handling 1M+ daily requests', style='List Bullet 2')
doc.add_paragraph('Mentored team of 5 junior developers', style='List Bullet 2')
doc.add_paragraph('Reduced API response time by 40% through optimization\n', style='List Bullet 2')

doc.add_paragraph('Full-Stack Developer | StartupXYZ | 2018-2021', style='List Bullet')
doc.add_paragraph('Built and maintained React-based frontend and Python/Django backend', style='List Bullet 2')
doc.add_paragraph('Implemented CI/CD pipeline reducing deployment time by 60%', style='List Bullet 2')
doc.add_paragraph('Designed and implemented RESTful APIs\n', style='List Bullet 2')

doc.add_heading('SKILLS', level=1)
doc.add_paragraph('Languages: Python, JavaScript, Java, SQL, HTML/CSS')
doc.add_paragraph('Frameworks: Django, React, FastAPI, Node.js')
doc.add_paragraph('Databases: PostgreSQL, MongoDB, Redis')
doc.add_paragraph('Tools: Git, Docker, Kubernetes, AWS\n')

doc.add_heading('EDUCATION', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science | State University | 2016, GPA: 3.8/4.0\n')

doc.add_heading('CERTIFICATIONS', level=1)
doc.add_paragraph('AWS Solutions Architect Associate (2022)', style='List Bullet')
doc.add_paragraph('Kubernetes Administrator (CKA) (2021)', style='List Bullet')

doc.save(test_resume_docx)

print("=" * 70)
print("RESUME PARSING TEST - With Searchable Document")
print("=" * 70)

# Parse the test resume
try:
    result = parse_resume(test_resume_docx)
    
    print("\n✅ PARSING SUCCESSFUL\n")
    print(f"Name: {result.get('name')}")
    print(f"Email: {result.get('email')}")
    print(f"Phone: {result.get('phone')}")
    print(f"Current Role: {result.get('current_role')}")
    print(f"Experience: {result.get('experience_years')} years")
    print(f"Education: {result.get('education')}")
    print(f"Skills: {result.get('skills')}")
    print(f"Red Flags: {result.get('red_flags') or 'None detected'}")
    print(f"Image-based: {result.get('is_image_based')}")
    
except Exception as e:
    print(f"\n❌ ERROR: {e}")

# Clean up
os.remove(test_resume_docx)

print("\n" + "=" * 70)
print("Demo Resume Status (Image-Based PDF)")
print("=" * 70)

# Now test the actual Demo Resume
demo_result = parse_resume("uploads/Demo Resume.pdf")
print(f"\n❌ Red Flags: {demo_result.get('red_flags')}")
print(f"Image-based: {demo_result.get('is_image_based')}")

print("\n" + "=" * 70)
print("WHAT YOU NEED TO DO")
print("=" * 70)
print("""
Your Demo Resume is an image-based PDF (scanned document).
The system detected this but cannot extract text because Tesseract-OCR
is not installed on this system.

CHOOSE ONE OF THESE SOLUTIONS:

✓ OPTION 1 - Convert PDF to searchable format (RECOMMENDED)
  1. Go to: https://smallpdf.com/pdf-to-pdf or similar tool
  2. Upload Demo Resume.pdf
  3. Download the searchable PDF
  4. Upload to HireIQ - it will parse correctly!

✓ OPTION 2 - Convert using Google Drive
  1. Upload Demo Resume.pdf to Google Drive
  2. Right-click → "Open with" → "Google Docs"
  3. Google Docs will OCR it
  4. File → Download → PDF
  5. Upload to HireIQ - it will parse correctly!

✓ OPTION 3 - Use an already searchable resume
  Try uploading resume_selected_1.pdf or resume_reject_1.pdf from the
  uploads folder - these might already be searchable

After fixing the resume, re-upload it to HireIQ and you'll see:
- Candidate name, email, and phone properly extracted
- Experience level correctly identified
- Skills automatically parsed
- AI-powered scoring and decision making in action
""")

